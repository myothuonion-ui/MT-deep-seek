"""
MT Pentester Session Report Generator
Produces a DOCX penetration-test report from a session dict (as returned by
orchestrator.get_session_report()).

Designed to be import-safe: if python-docx is not installed the module loads
without error and generate_report() raises ImportError with a clear message,
so the FastAPI server still starts even on minimal installs.

Usage (from backend):
    from core.report_generator import generate_report
    path = generate_report(session_report_dict, output_path="/tmp/report.docx")
    # → returns the path on success, raises on failure
"""

import json
import logging
import os
from datetime import datetime
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)

def _display_secret(value) -> str:
    raw = str(value or "")
    if os.getenv("INCLUDE_SECRETS_IN_REPORTS", "false").lower() in {"1", "true", "yes", "on"}:
        return raw
    return "********" if raw else ""

def _secure_file(path: str) -> None:
    try:
        os.chmod(path, 0o600)
    except OSError:
        pass

# ---------------------------------------------------------------------------
# Risk-level → colour mapping (Word RGB)
# ---------------------------------------------------------------------------
_RISK_COLORS = {
    "high":    (0xD3, 0x2F, 0x2F),  # red
    "medium":  (0xF5, 0x7F, 0x17),  # amber
    "low":     (0x2E, 0x7D, 0x32),  # green
    "unknown": (0x55, 0x55, 0x55),  # grey
}

_HEADER_BG = (0x1A, 0x23, 0x7E)   # dark indigo — title bar
_ACCENT_BG  = (0x37, 0x47, 0x4F)  # blue-grey — section headers


def _require_docx():
    try:
        from docx import Document                          # noqa: F401
        from docx.shared import Pt, RGBColor, Inches     # noqa: F401
        from docx.enum.text import WD_ALIGN_PARAGRAPH     # noqa: F401
        from docx.oxml.ns import qn                       # noqa: F401
        from docx.oxml import OxmlElement                 # noqa: F401
    except ImportError:
        raise ImportError(
            "python-docx is required for report generation. "
            "Install it with: pip install python-docx"
        )


def _set_cell_bg(cell, rgb: tuple):
    """Apply a solid background colour to a table cell (OOXML shading)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = "{:02X}{:02X}{:02X}".format(*rgb)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_section_heading(doc, text: str):
    """Add a styled section heading paragraph."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(*_ACCENT_BG)
    # Bottom border as a visual rule
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*_ACCENT_BG))
    pBdr.append(bottom)
    pPr.append(pBdr)
    doc.add_paragraph("")   # spacer


def _add_table_header_row(table, headers: List[str], widths_dxa: List[int]):
    """Write column headers into row 0 with dark background + white bold text."""
    from docx.shared import Pt, RGBColor
    row = table.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, widths_dxa)):
        cell = row.cells[i]
        cell.width = w
        _set_cell_bg(cell, _HEADER_BG)
        para = cell.paragraphs[0]
        run = para.add_run(hdr)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def generate_report(session_report: Dict, output_path: Optional[str] = None) -> str:
    """Generate a DOCX penetration-test report from a session report dict.

    Args:
        session_report: The dict returned by Orchestrator.get_session_report()
        output_path: Where to write the .docx. Defaults to /tmp/mt_report_{session_id}.docx

    Returns:
        Absolute path to the generated file.

    Raises:
        ImportError: if python-docx is not installed
        ValueError: if session_report is empty/invalid
    """
    _require_docx()

    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    session = session_report.get("session", {})
    session_id = session.get("session_id", "unknown")
    target_ip = session.get("target_ip", "N/A")
    target_domain = session.get("target_domain") or ""
    created_at = (session.get("created_at") or "")[:19].replace("T", " ")
    status = session.get("status", "unknown")
    stage = session.get("current_stage", "unknown")
    report_date = datetime.now().strftime("%Y-%m-%d %H:%M UTC")

    discovered_services: List[Dict] = session_report.get("discovered_services", [])
    discovered_hosts: List[Dict]    = session_report.get("discovered_hosts", [])
    vulnerabilities: List[Dict]     = session_report.get("vulnerabilities", [])
    commands: List[Dict]            = session_report.get("commands_executed", [])
    credentials: List[Dict]         = session_report.get("credentials", [])
    ai_decisions: List[Dict]        = session_report.get("ai_decisions", [])
    compromises: List[Dict]         = session.get("compromise_evidence", []) or []
    strategic_plan: List[Dict]      = session.get("strategic_plan", []) or []
    operator_instructions: List[str] = session.get("operator_instructions", []) or []
    reflections: List[str]          = session.get("reflections", []) or []

    if not output_path:
        output_path = f"/tmp/mt_report_{session_id[:12]}.docx"

    doc = Document()

    # --- Page margins (narrow) ---
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ============================================================
    # COVER PAGE
    # ============================================================
    doc.add_paragraph("")
    doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("MT Pentester")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = RGBColor(*_HEADER_BG)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Penetration Test Session Report")
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(*_ACCENT_BG)

    doc.add_paragraph("")

    meta_lines = [
        ("Target", f"{target_ip}" + (f" / {target_domain}" if target_domain else "")),
        ("Session ID", session_id),
        ("Session Created", created_at),
        ("Report Generated", report_date),
        ("Final Status", status.upper()),
        ("Final Stage", stage),
    ]
    meta_table = doc.add_table(rows=len(meta_lines), cols=2)
    meta_table.style = "Table Grid"
    total_w = int(section.page_width - section.left_margin - section.right_margin)
    label_w = int(total_w * 0.35)
    value_w = total_w - label_w
    for i, (label, value) in enumerate(meta_lines):
        row = meta_table.rows[i]
        lc = row.cells[0]
        vc = row.cells[1]
        lc.width = label_w
        vc.width = value_w
        _set_cell_bg(lc, _ACCENT_BG)
        lr = lc.paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        vr = vc.paragraphs[0].add_run(str(value))
        vr.font.size = Pt(9)

    doc.add_page_break()

    # ============================================================
    # 1. EXECUTIVE SUMMARY
    # ============================================================
    _add_section_heading(doc, "1. Executive Summary")

    high_v   = sum(1 for v in vulnerabilities if v.get("risk_level") == "high")
    medium_v = sum(1 for v in vulnerabilities if v.get("risk_level") == "medium")
    low_v    = sum(1 for v in vulnerabilities if v.get("risk_level") == "low")
    unverif  = sum(1 for v in vulnerabilities if v.get("status") == "unverified")

    summary_text = (
        f"This report summarises the results of an AI-directed penetration test session "
        f"targeting {target_ip}" + (f" ({target_domain})" if target_domain else "") + ". "
        f"The session discovered {len(discovered_hosts)} host(s) and "
        f"{len(discovered_services)} service(s). "
        f"A total of {len(vulnerabilities)} vulnerability finding(s) were recorded: "
        f"{high_v} high, {medium_v} medium, {low_v} low severity "
        f"(plus {unverif} unverified leads from web research). "
        f"{len(credentials)} credential(s) were captured. "
        f"{len(commands)} command(s) were executed during the session."
    )
    p = doc.add_paragraph(summary_text)
    p.style.font.size = Pt(10)

    doc.add_paragraph("")

    # Metrics summary table (1 row × 5 cols)
    metrics = [
        ("Hosts", str(len(discovered_hosts))),
        ("Services", str(len(discovered_services))),
        ("High Vulns", str(high_v)),
        ("Medium Vulns", str(medium_v)),
        ("Credentials", str(len(credentials))),
    ]
    mt = doc.add_table(rows=2, cols=len(metrics))
    mt.style = "Table Grid"
    col_w = total_w // len(metrics)
    for i, (label, val) in enumerate(metrics):
        hcell = mt.rows[0].cells[i]
        vcell = mt.rows[1].cells[i]
        hcell.width = col_w
        vcell.width = col_w
        _set_cell_bg(hcell, _ACCENT_BG)
        hr2 = hcell.paragraphs[0].add_run(label)
        hr2.bold = True
        hr2.font.size = Pt(8)
        hr2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr2 = vcell.paragraphs[0].add_run(val)
        vr2.bold = True
        vr2.font.size = Pt(14)
        vcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # ============================================================
    # 2. DISCOVERED SERVICES
    # ============================================================
    _add_section_heading(doc, "2. Discovered Services")

    if discovered_services:
        svc_headers = ["Host", "Port", "Service", "Version", "State"]
        svc_widths  = [int(total_w * p) for p in [0.20, 0.10, 0.20, 0.35, 0.15]]
        st2 = doc.add_table(rows=1 + len(discovered_services), cols=5)
        st2.style = "Table Grid"
        _add_table_header_row(st2, svc_headers, svc_widths)
        for i, svc in enumerate(discovered_services):
            row = st2.rows[i + 1]
            vals = [
                str(svc.get("host") or target_ip),
                str(svc.get("port") or ""),
                str(svc.get("service") or ""),
                str(svc.get("version") or ""),
                str(svc.get("state") or "open"),
            ]
            for j, (cell, val, w) in enumerate(zip(row.cells, vals, svc_widths)):
                cell.width = w
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9)
    else:
        doc.add_paragraph("No services discovered in this session.")

    doc.add_paragraph("")

    # ============================================================
    # 3. VULNERABILITY FINDINGS
    # ============================================================
    _add_section_heading(doc, "3. Vulnerability Findings")

    if vulnerabilities:
        # Group by risk descending
        ordered = sorted(
            vulnerabilities,
            key=lambda v: {"high": 0, "medium": 1, "low": 2}.get(v.get("risk_level", ""), 3)
        )
        vheaders = ["#", "Risk", "Name", "Host:Port", "CVE(s)", "Source", "Status"]
        vwidths  = [int(total_w * p) for p in [0.04, 0.07, 0.24, 0.15, 0.18, 0.14, 0.13]]
        # Adjust last to fill remainder
        vwidths[-1] = total_w - sum(vwidths[:-1])

        vt = doc.add_table(rows=1 + len(ordered), cols=7)
        vt.style = "Table Grid"
        _add_table_header_row(vt, vheaders, vwidths)

        for i, vuln in enumerate(ordered):
            row = vt.rows[i + 1]
            risk = (vuln.get("risk_level") or "unknown").lower()
            rgb = _RISK_COLORS.get(risk, _RISK_COLORS["unknown"])
            cves = ", ".join(vuln.get("cve_ids") or []) or "—"
            host_port = f"{vuln.get('host') or ''}"
            if vuln.get("port"):
                host_port += f":{vuln['port']}"
            vals = [
                str(i + 1),
                risk.upper(),
                (vuln.get("name") or "")[:80],
                host_port,
                cves[:60],
                (vuln.get("source_tool") or "")[:20],
                (vuln.get("status") or "confirmed")[:15],
            ]
            for j, (cell, val, w) in enumerate(zip(row.cells, vals, vwidths)):
                cell.width = w
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(8)
                if j == 1:  # Risk column - colour text
                    run.bold = True
                    run.font.color.rgb = RGBColor(*rgb)

        # Per-finding details (expandable detail block for high/medium)
        doc.add_paragraph("")
        _add_section_heading(doc, "3.1  Finding Details")
        for i, vuln in enumerate(ordered):
            if vuln.get("risk_level") not in ("high", "medium"):
                continue
            p = doc.add_paragraph()
            r = p.add_run(f"[{i+1}] {vuln.get('name') or 'Unnamed'}")
            r.bold = True
            r.font.size = Pt(10)
            risk = (vuln.get("risk_level") or "unknown").lower()
            r.font.color.rgb = RGBColor(*_RISK_COLORS.get(risk, _RISK_COLORS["unknown"]))

            for label, key in [
                ("Description", "description"), ("Affected Software", "service_version"),
            ]:
                val = vuln.get(key) or ""
                if val:
                    lp = doc.add_paragraph()
                    lr = lp.add_run(f"{label}: ")
                    lr.bold = True
                    lr.font.size = Pt(9)
                    lp.add_run(val[:500]).font.size = Pt(9)

            cves = ", ".join(vuln.get("cve_ids") or [])
            if cves:
                lp = doc.add_paragraph()
                lr = lp.add_run("CVE(s): ")
                lr.bold = True
                lr.font.size = Pt(9)
                lp.add_run(cves).font.size = Pt(9)

            refs = vuln.get("reference_urls") or []
            if refs:
                lp = doc.add_paragraph()
                lr = lp.add_run("References: ")
                lr.bold = True
                lr.font.size = Pt(9)
                lp.add_run("; ".join(refs[:3])).font.size = Pt(9)

            doc.add_paragraph("")
    else:
        doc.add_paragraph("No vulnerabilities recorded in this session.")

    doc.add_paragraph("")

    # ============================================================
    # 4. CREDENTIALS CAPTURED
    # ============================================================
    _add_section_heading(doc, "4. Credentials Captured")

    if credentials:
        cheaders = ["Username", "Secret", "Type", "Service", "Discovered"]
        cwidths  = [int(total_w * p) for p in [0.18, 0.30, 0.10, 0.12, 0.20]]
        cwidths[-1] = total_w - sum(cwidths[:-1])
        ct = doc.add_table(rows=1 + len(credentials), cols=5)
        ct.style = "Table Grid"
        _add_table_header_row(ct, cheaders, cwidths)
        for i, cred in enumerate(credentials):
            row = ct.rows[i + 1]
            vals = [
                cred.get("username") or "",
                _display_secret(cred.get("secret")),
                cred.get("secret_type") or "password",
                cred.get("service") or "",
                (cred.get("discovered_at") or "")[:19],
            ]
            for cell, val, w in zip(row.cells, vals, cwidths):
                cell.width = w
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9)
    else:
        doc.add_paragraph("No credentials were captured during this session.")

    doc.add_paragraph("")

    # ============================================================
    # 4a. CONFIRMED COMPROMISES
    # ============================================================
    _add_section_heading(doc, "4.1  Confirmed Compromises")
    if compromises:
        cmp_headers = ["#", "Service:Port", "Host", "Privilege", "Via Command", "Signal"]
        cmp_widths  = [int(total_w * p) for p in [0.05, 0.16, 0.15, 0.14, 0.34, 0.16]]
        cmp_widths[-1] = total_w - sum(cmp_widths[:-1])
        ctab = doc.add_table(rows=1 + len(compromises), cols=6)
        ctab.style = "Table Grid"
        _add_table_header_row(ctab, cmp_headers, cmp_widths)
        for i, comp in enumerate(compromises):
            row = ctab.rows[i + 1]
            vals = [
                str(i + 1),
                f"{comp.get('service','?')}:{comp.get('port','?')}",
                str(comp.get("host") or ""),
                str(comp.get("privilege") or "?"),
                (comp.get("command") or "")[:60],
                (comp.get("signal") or "")[:20],
            ]
            for j, (cell, val, w) in enumerate(zip(row.cells, vals, cmp_widths)):
                cell.width = w
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(8)
                if j == 3 and "root" in str(comp.get("privilege", "")).lower():
                    run.bold = True
                    run.font.color.rgb = RGBColor(*_RISK_COLORS["high"])
        # Proof snippets
        for i, comp in enumerate(compromises):
            proof = (comp.get("proof") or "").strip()
            if not proof:
                continue
            pp = doc.add_paragraph()
            pr = pp.add_run(f"[{i+1}] proof ({comp.get('service','?')}:{comp.get('port','?')}):")
            pr.bold = True
            pr.font.size = Pt(8)
            op = doc.add_paragraph()
            op.paragraph_format.left_indent = Inches(0.3)
            orr = op.add_run(proof[:500])
            orr.font.name = "Courier New"
            orr.font.size = Pt(7)
    else:
        doc.add_paragraph("No confirmed code-execution / shell access was recorded.")

    doc.add_paragraph("")

    # ============================================================
    # 4b. ATTACK PLAN & OPERATOR STEERING
    # ============================================================
    _add_section_heading(doc, "4.2  Attack Plan & Operator Steering")
    if strategic_plan:
        pp = doc.add_paragraph()
        pp.add_run("Strategic plan (AI planner):").bold = True
        for step in strategic_plan:
            sp = doc.add_paragraph(style="List Bullet")
            sp.add_run(f"[{step.get('status','pending')}] {step.get('step','')}").font.size = Pt(9)
    if operator_instructions:
        pp = doc.add_paragraph()
        pp.add_run("Operator steering instructions (chronological):").bold = True
        for instr in operator_instructions:
            ip = doc.add_paragraph(style="List Bullet")
            ip.add_run(str(instr)).font.size = Pt(9)
    if reflections:
        pp = doc.add_paragraph()
        pp.add_run("Latest strategist reflections:").bold = True
        for refl in reflections[-3:]:
            rp = doc.add_paragraph(style="List Bullet")
            rp.add_run(str(refl)[:400]).font.size = Pt(9)
    if not (strategic_plan or operator_instructions or reflections):
        doc.add_paragraph("No strategic plan or operator instructions recorded.")

    doc.add_paragraph("")

    # ============================================================
    # 5. EXECUTED COMMANDS LOG
    # ============================================================
    _add_section_heading(doc, "5. Executed Commands Log")

    if commands:
        for i, cmd in enumerate(commands):
            p = doc.add_paragraph()
            r = p.add_run(f"[{i+1}]  {cmd.get('command','')[:120]}")
            r.font.name = "Courier New"
            r.font.size = Pt(8)
            r.bold = True

            ts = (cmd.get("timestamp") or "")[:19]
            ok = "✓" if cmd.get("success") else "✗"
            meta = doc.add_paragraph()
            mr = meta.add_run(f"     {ok}  {ts}")
            mr.font.size = Pt(8)
            mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            output = (cmd.get("output") or "")[:600]
            if output:
                op = doc.add_paragraph()
                op.paragraph_format.left_indent = Inches(0.3)
                orr = op.add_run(output)
                orr.font.name = "Courier New"
                orr.font.size = Pt(7)
                orr.font.color.rgb = RGBColor(0x00, 0x33, 0x00)

            doc.add_paragraph("")
    else:
        doc.add_paragraph("No commands executed in this session.")

    # ============================================================
    # 6. ATTACK DECISIONS & IDEAS (full, chronological)
    # ============================================================
    if ai_decisions:
        doc.add_page_break()
        _add_section_heading(
            doc, f"6. Attack Decisions & Ideas — chronological ({len(ai_decisions)})"
        )
        _CTX_TAG = {
            "auto_pivot": "PIVOT", "pivot_limit_reached": "PIVOT-LIMIT",
            "loop_prevention": "LOOP", "loop_error": "ERROR",
            "no_next_step": "NO-STEP", "watchdog_stalled": "WATCHDOG",
            "operator_instruction": "OPERATOR", "handler_started": "LISTENER",
            "shell_caught": "SHELL", "self_critique_reject": "VETOED",
        }
        for idx, decision in enumerate(ai_decisions, 1):
            ctx = decision.get("context", "")
            tag = _CTX_TAG.get(ctx, "")
            ts = (decision.get("timestamp") or "")[:19].replace("T", " ")
            p = doc.add_paragraph()
            head = f"[{idx}] {ts}  |  Phase: {decision.get('attack_phase', ctx or '?')}"
            if decision.get("risk_level"):
                head += f"  |  Risk: {decision.get('risk_level')}"
            if tag:
                head += f"  |  {tag}"
            tr2 = p.add_run(head)
            tr2.bold = True
            tr2.font.size = Pt(9)
            if tag in ("SHELL", "OPERATOR"):
                tr2.font.color.rgb = RGBColor(*_RISK_COLORS["high"])

            cmd = (decision.get("suggested_command") or "").strip()
            if cmd:
                cmd_p = doc.add_paragraph()
                cmd_p.paragraph_format.left_indent = Inches(0.3)
                cr = cmd_p.add_run(f"$ {cmd[:160]}")
                cr.font.name = "Courier New"
                cr.font.size = Pt(8)

            reasoning = (decision.get("reasoning") or "")[:600]
            if reasoning:
                rp = doc.add_paragraph()
                rp.paragraph_format.left_indent = Inches(0.3)
                rr = rp.add_run(reasoning)
                rr.font.size = Pt(8)
                rr.font.color.rgb = RGBColor(0x33, 0x33, 0x55)

    # ============================================================
    # DISCLAIMER FOOTER
    # ============================================================
    doc.add_page_break()
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = disc.add_run("LEGAL DISCLAIMER")
    dr.bold = True
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(*_RISK_COLORS["high"])

    doc.add_paragraph(
        "This report was generated by MT Pentester, an AI-assisted penetration testing "
        "framework. All testing activity recorded in this report was performed only against "
        "systems for which explicit written authorisation was obtained prior to testing "
        "(as confirmed by the authorization_confirmed flag in session metadata). "
        "Vulnerability findings derived from unverified web research (source: threat-intel-cache) "
        "are marked as such and must be independently corroborated before being treated as "
        "confirmed. The operator is solely responsible for the legality and scope of all "
        "testing activity."
    ).runs[0].font.size = Pt(9)

    doc.save(output_path)
    logger.info(f"Report saved to {output_path}")
    _secure_file(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Markdown report generator — pure Python, ZERO dependencies. Always available
# (no python-docx / fpdf2 needed), so it's the reliable fallback format. Renders
# EVERYTHING the session captured, in order: findings, all executed commands,
# and every attack decision/idea chronologically.
# ---------------------------------------------------------------------------

def generate_markdown_report(session_report: Dict, output_path: Optional[str] = None) -> str:
    """Build a complete Markdown pentest report from a session report dict.

    Includes (in order): metadata, executive summary, discovered services,
    vulnerability findings, credentials, confirmed compromises, attack plan +
    operator steering, the FULL executed-command log (chronological, with output),
    and the FULL attack-decision/idea log (chronological). No third-party deps.
    """
    session = session_report.get("session", {})
    sid = session.get("session_id", "unknown")
    target_ip = session.get("target_ip", "N/A")
    target_domain = session.get("target_domain") or ""
    created = (session.get("created_at") or "")[:19].replace("T", " ")
    status = session.get("status", "unknown")
    stage = session.get("current_stage", "unknown")
    gen = datetime.now().strftime("%Y-%m-%d %H:%M UTC")

    services = session_report.get("discovered_services", []) or []
    hosts = session_report.get("discovered_hosts", []) or []
    vulns = session_report.get("vulnerabilities", []) or []
    commands = session_report.get("commands_executed", []) or []
    creds = session_report.get("credentials", []) or []
    decisions = session_report.get("ai_decisions", []) or []
    compromises = session.get("compromise_evidence", []) or []
    plan = session.get("strategic_plan", []) or []
    operator_instructions = session.get("operator_instructions", []) or []
    reflections = session.get("reflections", []) or []
    exhausted = session.get("exhausted_services", []) or []

    high = sum(1 for v in vulns if v.get("risk_level") == "high")
    med = sum(1 for v in vulns if v.get("risk_level") == "medium")
    low = sum(1 for v in vulns if v.get("risk_level") == "low")

    def esc(x) -> str:
        # Keep table cells intact: strip newlines and escape pipes.
        return str(x if x is not None else "").replace("|", "\\|").replace("\n", " ").strip()

    L: List[str] = []
    a = L.append

    a(f"# MT Pentester — Penetration Test Report")
    a("")
    a(f"**Target:** `{target_ip}`" + (f" / `{target_domain}`" if target_domain else ""))
    a(f"**Session ID:** `{sid}`  ")
    a(f"**Created:** {created}  ")
    a(f"**Report generated:** {gen}  ")
    a(f"**Final status:** {status}  **Final stage:** {stage}")
    a("")
    a("---")
    a("")

    # 1. Executive summary
    a("## 1. Executive Summary")
    a("")
    a(f"An AI-directed penetration test of **{target_ip}**"
      + (f" ({target_domain})" if target_domain else "")
      + f" discovered **{len(hosts)} host(s)** and **{len(services)} service(s)**. "
      + f"**{len(vulns)} vulnerability finding(s)** were recorded "
      + f"({high} high, {med} medium, {low} low), "
      + f"**{len(creds)} credential(s)** captured, "
      + f"**{len(compromises)} confirmed compromise(s)**, "
      + f"across **{len(commands)} executed command(s)** and "
      + f"**{len(decisions)} AI decision(s)**.")
    a("")

    # 2. Services
    a("## 2. Discovered Services")
    a("")
    if services:
        a("| Host | Port | Service | Version | State |")
        a("|------|------|---------|---------|-------|")
        for s in services:
            a(f"| {esc(s.get('host') or target_ip)} | {esc(s.get('port'))} | "
              f"{esc(s.get('service'))} | {esc(s.get('version'))} | "
              f"{esc(s.get('state') or s.get('test_state') or 'open')} |")
    else:
        a("_No services discovered._")
    a("")

    # 3. Vulnerabilities
    a("## 3. Vulnerability Findings")
    a("")
    if vulns:
        ordered = sorted(vulns, key=lambda v: {"high": 0, "medium": 1, "low": 2}
                         .get(v.get("risk_level", ""), 3))
        a("| # | Risk | Name | Host:Port | CVE(s) | Source | Status |")
        a("|---|------|------|-----------|--------|--------|--------|")
        for i, v in enumerate(ordered, 1):
            cids = v.get("cve_ids") or []
            if isinstance(cids, str):
                cids = [cids]
            hp = esc(v.get("host") or "")
            if v.get("port"):
                hp += f":{v.get('port')}"
            a(f"| {i} | {esc((v.get('risk_level') or 'unknown').upper())} | "
              f"{esc(v.get('name'))} | {hp} | {esc(', '.join(cids) or '—')} | "
              f"{esc(v.get('source_tool'))} | {esc(v.get('status') or 'confirmed')} |")
        a("")
        # Details for high/medium
        detail = [v for v in ordered if v.get("risk_level") in ("high", "medium")]
        if detail:
            a("### 3.1 Finding Details")
            a("")
            for i, v in enumerate(detail, 1):
                a(f"**[{i}] {v.get('name') or 'Unnamed'}** "
                  f"({(v.get('risk_level') or '').upper()})")
                if v.get("description"):
                    a(f"- Description: {esc(v.get('description'))[:600]}")
                if v.get("service_version"):
                    a(f"- Affected: {esc(v.get('service_version'))}")
                cids = v.get("cve_ids") or []
                if cids:
                    a(f"- CVE(s): {esc(', '.join(cids if isinstance(cids, list) else [cids]))}")
                refs = v.get("reference_urls") or []
                if refs:
                    a(f"- References: {esc('; '.join(refs[:3]))}")
                a("")
    else:
        a("_No vulnerabilities recorded._")
    a("")

    # 4. Credentials
    a("## 4. Credentials Captured")
    a("")
    if creds:
        a("| Username | Secret | Type | Service | Discovered |")
        a("|----------|--------|------|---------|------------|")
        for c in creds:
            a(f"| {esc(c.get('username'))} | {esc(_display_secret(c.get('secret')))} | "
              f"{esc(c.get('secret_type') or 'password')} | {esc(c.get('service'))} | "
              f"{esc((c.get('discovered_at') or '')[:19])} |")
    else:
        a("_No credentials captured._")
    a("")

    # 4.1 Confirmed compromises
    a("## 4.1 Confirmed Compromises")
    a("")
    if compromises:
        a("| # | Service:Port | Host | Privilege | Via | Signal |")
        a("|---|--------------|------|-----------|-----|--------|")
        for i, c in enumerate(compromises, 1):
            a(f"| {i} | {esc(c.get('service'))}:{esc(c.get('port'))} | {esc(c.get('host'))} | "
              f"**{esc(c.get('privilege'))}** | {esc((c.get('command') or '')[:70])} | "
              f"{esc(c.get('signal'))} |")
        a("")
        for i, c in enumerate(compromises, 1):
            if c.get("proof"):
                a(f"Proof [{i}] — {esc(c.get('service'))}:{esc(c.get('port'))}:")
                a("```")
                a(str(c.get("proof"))[:600])
                a("```")
    else:
        a("_No confirmed code-execution / shell access recorded._")
    a("")

    # 4.2 Attack plan & operator steering
    a("## 4.2 Attack Plan & Operator Steering")
    a("")
    if plan:
        a("**Strategic plan (AI planner):**")
        for step in plan:
            a(f"- [{esc(step.get('status') or 'pending')}] {esc(step.get('step'))}")
        a("")
    if operator_instructions:
        a("**Operator steering instructions (chronological):**")
        for ins in operator_instructions:
            a(f"- {esc(ins)}")
        a("")
    if reflections:
        a("**Latest strategist reflections:**")
        for r in reflections[-3:]:
            a(f"- {esc(r)[:400]}")
        a("")
    if exhausted:
        a(f"**Exhausted vectors:** {esc(', '.join(exhausted))}")
        a("")
    if not (plan or operator_instructions or reflections or exhausted):
        a("_No strategic plan or operator instructions recorded._")
        a("")

    # 5. Executed commands (full, chronological)
    a(f"## 5. Executed Commands Log ({len(commands)})")
    a("")
    if commands:
        for i, cmd in enumerate(commands, 1):
            ok = "✓" if cmd.get("success") else "✗"
            ts = (cmd.get("timestamp") or "")[:19].replace("T", " ")
            a(f"**[{i}] {ok} `{esc(cmd.get('command'))}`**  \n_{ts}_")
            out = (cmd.get("output") or "").strip()
            if out:
                a("```")
                a(out[:1500])
                a("```")
            a("")
    else:
        a("_No commands executed._")
    a("")

    # 6. Attack decisions & ideas (full, chronological)
    a(f"## 6. Attack Decisions & Ideas — chronological ({len(decisions)})")
    a("")
    _CTX_TAG = {
        "auto_pivot": "PIVOT", "pivot_limit_reached": "PIVOT-LIMIT",
        "loop_prevention": "LOOP", "loop_error": "ERROR", "no_next_step": "NO-STEP",
        "watchdog_stalled": "WATCHDOG", "operator_instruction": "OPERATOR",
        "handler_started": "LISTENER", "shell_caught": "SHELL",
        "self_critique_reject": "VETOED",
    }
    if decisions:
        for i, d in enumerate(decisions, 1):
            ctx = d.get("context", "")
            tag = _CTX_TAG.get(ctx, "")
            ts = (d.get("timestamp") or "")[:19].replace("T", " ")
            head = (f"**[{i}] {ts} · {esc(d.get('attack_phase', ctx or '?'))}**"
                    + (f" · _{esc(d.get('risk_level'))}_" if d.get("risk_level") else "")
                    + (f" · **{tag}**" if tag else ""))
            a(head)
            cmd = (d.get("suggested_command") or "").strip()
            if cmd:
                a(f"- Suggested: `{esc(cmd)}`")
            reason = (d.get("reasoning") or "").strip()
            if reason:
                a(f"- Reasoning: {esc(reason)[:800]}")
            a("")
    else:
        a("_No AI decisions recorded._")
    a("")

    a("---")
    a("")
    a("> **Legal disclaimer.** Generated by MT Pentester. All activity herein was "
      "performed only against systems with explicit prior written authorisation. "
      "Findings from unverified web research must be independently corroborated. "
      "The operator is solely responsible for the legality and scope of all testing.")
    a("")

    if not output_path:
        output_path = f"/tmp/mt_report_{sid[:12]}.md"
    with open(output_path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(L))
    logger.info(f"Markdown report saved to {output_path}")
    _secure_file(output_path)
    return output_path


# ---------------------------------------------------------------------------
# PDF report generator (uses fpdf2 — pure Python, no LibreOffice needed)
# ---------------------------------------------------------------------------

def _require_fpdf():
    try:
        from fpdf import FPDF  # noqa: F401
    except ImportError:
        raise ImportError(
            "fpdf2 is required for PDF report generation. "
            "Install it with: pip install fpdf2"
        )


def generate_pdf_report(session_report: Dict, output_path: Optional[str] = None) -> str:
    """Generate a PDF penetration-test report.

    Args:
        session_report: dict as returned by orchestrator.get_session_report()
        output_path: where to write the file. Defaults to /tmp/mt_report_<id>.pdf

    Returns:
        Absolute path to the generated PDF.
    """
    _require_fpdf()
    from fpdf import FPDF

    meta = session_report.get("session", {})
    session_id = meta.get("session_id", "unknown")[:12]
    target = meta.get("target_ip", "—")
    created = meta.get("created_at", "—")
    status = meta.get("status", "—")
    generated_at = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")

    vulns: List[Dict] = session_report.get("vulnerabilities", [])
    # NOTE: get_session_report() uses keys 'commands_executed' and
    # 'discovered_services' — the old 'commands'/'services' keys were wrong and
    # silently produced empty sections.
    commands: List[Dict] = session_report.get("commands_executed", []) or session_report.get("commands", [])
    services: List[Dict] = session_report.get("discovered_services", []) or session_report.get("services", [])
    credentials: List[Dict] = session_report.get("credentials", [])
    ai_decisions: List[Dict] = session_report.get("ai_decisions", [])
    compromises: List[Dict] = meta.get("compromise_evidence", []) or []
    operator_instructions: List[str] = meta.get("operator_instructions", []) or []

    # Risk summary
    high_c = sum(1 for v in vulns if v.get("risk_level") == "high")
    med_c  = sum(1 for v in vulns if v.get("risk_level") == "medium")
    low_c  = sum(1 for v in vulns if v.get("risk_level") == "low")

    if output_path is None:
        output_path = os.path.join("/tmp", f"mt_report_{session_id}.pdf")

    # ── PDF setup ──────────────────────────────────────────────────────────
    class _PDF(FPDF):
        def header(self):
            self.set_fill_color(26, 35, 126)       # dark indigo
            self.rect(0, 0, 210, 14, "F")
            self.set_font("Helvetica", "B", 9)
            self.set_text_color(255, 255, 255)
            self.set_xy(8, 3)
            self.cell(0, 8, "MT Pentester  |  Penetration Test Report  — CONFIDENTIAL")
            self.set_text_color(0, 0, 0)

        def footer(self):
            self.set_y(-12)
            self.set_font("Helvetica", "", 7)
            self.set_text_color(120, 120, 120)
            self.cell(0, 6, f"Page {self.page_no()}", align="C")
            self.set_text_color(0, 0, 0)

    pdf = _PDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=16)
    pdf.add_page()
    pdf.set_margins(14, 18, 14)

    def _h1(txt: str):
        pdf.ln(4)
        pdf.set_fill_color(55, 71, 79)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 8, f"  {txt}", fill=True, ln=True)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(2)

    def _row(label: str, value: str, bold_val: bool = False):
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(50, 6, label + ":", ln=False)
        pdf.set_font("Helvetica", "B" if bold_val else "", 9)
        pdf.multi_cell(0, 6, str(value))

    def _para(txt: str, size: int = 9):
        pdf.set_font("Helvetica", "", size)
        pdf.multi_cell(0, 5, txt)
        pdf.ln(1)

    # ── Cover / summary ───────────────────────────────────────────────────
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(26, 35, 126)
    pdf.cell(0, 12, "MT Pentester", ln=True, align="C")
    pdf.set_font("Helvetica", "", 12)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 7, "AI-Driven Penetration Test Report", ln=True, align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(6)

    _h1("Session Metadata")
    _row("Session ID", session_id)
    _row("Target", target)
    _row("Status", status)
    _row("Started", str(created))
    _row("Generated", generated_at)
    pdf.ln(2)

    _h1("Executive Summary")
    _row("Total Vulnerabilities", str(len(vulns)))
    _row("High Risk", str(high_c), bold_val=high_c > 0)
    _row("Medium Risk", str(med_c))
    _row("Low Risk", str(low_c))
    _row("Services Discovered", str(len(services)))
    _row("Commands Executed", str(len(commands)))
    _row("Credentials Captured", str(len(credentials)))
    pdf.ln(2)

    # ── Vulnerabilities ───────────────────────────────────────────────────
    if vulns:
        _h1(f"Vulnerability Findings ({len(vulns)})")
        for i, v in enumerate(
            sorted(vulns, key=lambda x: {"high": 0, "medium": 1, "low": 2}.get(x.get("risk_level", ""), 3)),
            start=1
        ):
            risk = v.get("risk_level", "unknown").upper()
            name = v.get("name", "Unknown")
            host = v.get("host", "—")
            port = v.get("port") or "—"
            svc  = v.get("service", "—")
            # cve_ids is already a list (get_session_report); tolerate a JSON
            # string too, for safety.
            _cids = v.get("cve_ids") or []
            if isinstance(_cids, str):
                try:
                    _cids = json.loads(_cids)
                except Exception:
                    _cids = [_cids] if _cids else []
            cves = ", ".join(_cids) or "—"
            desc = v.get("description", "")

            # Risk colour
            rc = {"HIGH": (211, 47, 47), "MEDIUM": (245, 127, 23), "LOW": (46, 125, 50)}.get(risk, (100, 100, 100))
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(*rc)
            pdf.cell(0, 6, f"[{risk}] {i}. {name}", ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", "", 8)
            pdf.cell(0, 5, f"Host: {host}  |  Port: {port}  |  Service: {svc}  |  CVEs: {cves}", ln=True)
            if desc:
                pdf.set_font("Helvetica", "I", 8)
                pdf.multi_cell(0, 5, desc[:300] + ("…" if len(desc) > 300 else ""))
            pdf.ln(1)

    # ── Services ─────────────────────────────────────────────────────────
    if services:
        _h1(f"Discovered Services ({len(services)})")
        pdf.set_font("Helvetica", "B", 8)
        # Header row
        col_w = [30, 20, 18, 60, 52]
        headers = ["Host", "Port", "State", "Service", "Version"]
        for w, h in zip(col_w, headers):
            pdf.cell(w, 6, h, border=1)
        pdf.ln()
        pdf.set_font("Helvetica", "", 8)
        for s in services[:40]:
            host = str(s.get("host", ""))[:18]
            port = str(s.get("port", ""))
            state = str(s.get("state", ""))
            svc  = str(s.get("service", ""))[:28]
            ver  = str(s.get("version", ""))[:30]
            for w, val in zip(col_w, [host, port, state, svc, ver]):
                pdf.cell(w, 5, val, border=1)
            pdf.ln()
        pdf.ln(2)

    # ── Credentials ───────────────────────────────────────────────────────
    if credentials:
        _h1(f"Captured Credentials ({len(credentials)})")
        pdf.set_font("Helvetica", "B", 8)
        for w, h in zip([40, 60, 25, 30, 25], ["Username", "Secret", "Type", "Service", "Host"]):
            pdf.cell(w, 6, h, border=1)
        pdf.ln()
        pdf.set_font("Helvetica", "", 8)
        for c in credentials:
            secret = _display_secret(c.get("secret", ""))
            if len(secret) > 28:
                secret = secret[:25] + "…"
            for w, val in zip([40, 60, 25, 30, 25], [
                str(c.get("username", ""))[:22],
                secret,
                str(c.get("secret_type", ""))[:10],
                str(c.get("service", ""))[:14],
                str(c.get("host", ""))[:14],
            ]):
                pdf.cell(w, 5, val, border=1)
            pdf.ln()
        pdf.ln(2)

    # ── Confirmed compromises ─────────────────────────────────────────────
    if compromises:
        _h1(f"Confirmed Compromises ({len(compromises)})")
        for i, comp in enumerate(compromises, start=1):
            priv = str(comp.get("privilege", "?"))
            rc = (211, 47, 47) if ("root" in priv.lower() or "system" in priv.lower()
                                   or "admin" in priv.lower()) else (245, 127, 23)
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(*rc)
            pdf.cell(0, 6, f"{i}. {comp.get('service','?')}:{comp.get('port','?')} "
                           f"on {comp.get('host','?')} - {priv}", ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", "", 8)
            pdf.multi_cell(0, 5, f"via: {str(comp.get('command',''))[:160]}")
            proof = str(comp.get("proof") or "").strip()
            if proof:
                pdf.set_font("Helvetica", "I", 7)
                pdf.multi_cell(0, 4, proof[:300])
            pdf.ln(1)

    # ── Operator steering ─────────────────────────────────────────────────
    if operator_instructions:
        _h1(f"Operator Steering Instructions ({len(operator_instructions)})")
        pdf.set_font("Helvetica", "", 8)
        for i, instr in enumerate(operator_instructions, start=1):
            pdf.multi_cell(0, 5, f"{i}. {str(instr)[:200]}")
        pdf.ln(1)

    # ── Commands (full, chronological) ────────────────────────────────────
    if commands:
        _h1(f"Commands Log ({len(commands)})")
        for i, cmd in enumerate(commands, start=1):
            ok = cmd.get("success", False)
            pdf.set_font("Helvetica", "B", 8)
            pdf.set_text_color(30, 100, 30) if ok else pdf.set_text_color(180, 30, 30)
            pdf.cell(0, 5, f"{i}. {'✓' if ok else '✗'}  {str(cmd.get('command', ''))[:100]}", ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", "", 7)
            out = str(cmd.get("output") or "").strip()[:400]
            if out:
                pdf.multi_cell(0, 4, out)
            pdf.ln(1)

    # ── Attack decisions & ideas (full, chronological) ────────────────────
    if ai_decisions:
        _h1(f"Attack Decisions & Ideas ({len(ai_decisions)})")
        for i, d in enumerate(ai_decisions, start=1):
            ctx = d.get("context", "")
            pdf.set_font("Helvetica", "B", 8)
            pdf.multi_cell(0, 5, f"{i}. [{(d.get('timestamp') or '')[:19]}] "
                                 f"{d.get('attack_phase', ctx or '?')}"
                                 + (f"  ({ctx})" if ctx else ""))
            cmd = str(d.get("suggested_command") or "").strip()
            if cmd:
                pdf.set_font("Helvetica", "", 7)
                pdf.multi_cell(0, 4, f"$ {cmd[:160]}")
            reason = str(d.get("reasoning") or "").strip()
            if reason:
                pdf.set_font("Helvetica", "I", 7)
                pdf.multi_cell(0, 4, reason[:300])
            pdf.ln(1)

    # ── Legal disclaimer ──────────────────────────────────────────────────
    pdf.add_page()
    _h1("Legal Disclaimer")
    _para(
        "This report was generated by MT Pentester and is intended SOLELY for use by "
        "authorised security professionals operating on systems for which explicit written "
        "authorisation was obtained prior to testing. Vulnerability findings derived from "
        "unverified web research are marked as such and must be independently corroborated. "
        "The operator is solely responsible for the legality and scope of all testing activity."
    )

    pdf.output(output_path)
    logger.info(f"PDF report saved to {output_path}")
    _secure_file(output_path)
    return output_path
